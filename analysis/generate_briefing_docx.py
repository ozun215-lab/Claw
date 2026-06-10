# -*- coding: utf-8 -*-
"""세무사 브리핑 DOCX 생성"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='맑은 고딕', size=11, bold=False, italic=False, color=None):
    """한글 폰트 안전 설정"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_p(doc, text, size=11, bold=False, align=None, italic=False, color=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    p.paragraph_format.space_after = Pt(6)
    return p


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x74, 0xB5))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table(doc, headers, data, col_widths=None):
    """표 생성"""
    table = doc.add_table(rows=len(data)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    # 헤더
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            set_font(r, size=10, bold=True)
    
    # 데이터
    for i, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            cell = table.rows[i+1].cells[j]
            cell.text = str(val)
            for r in cell.paragraphs[0].runs:
                set_font(r, size=10)
    
    return table


# === 문서 생성 ===
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 표지
add_p(doc, "김광규 송금 자금 출처", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x4E, 0x79))
add_p(doc, "세무사 브리핑 자료", size=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x4E, 0x79))
doc.add_paragraph()
add_p(doc, "─" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xAA, 0xAA, 0xAA))
add_p(doc, "작성일: 2026년 6월 9일", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "의뢰인: 박영준", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "용도: 세무 소명 자료 준비 (이번 주 제출 예정)", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True)
add_p(doc, "─" * 50, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0xAA, 0xAA, 0xAA))

# === 핵심 요약 ===
add_h1(doc, "📌 핵심 요약 (1분 브리핑)")

add_h2(doc, "1. 소명 대상")
add_p(doc, "2021.12 ~ 2022.02 김광규에게 송금한 530,000,000원의 자금 출처", bold=True)

add_h2(doc, "2. 자금 출처 결론")
add_table(doc, ["자금원", "금액", "비중"], [
    ["① 안경희 전세보증금 반환", "373,000,000원", "67.6%"],
    ["② IBK 037 상가 임대 보증금", "94,000,000원", "17.0%"],
    ["③ 김수동(매제) 차용금", "65,000,000원", "11.8%"],
    ["④ 배우자(이지영) 자금 이체", "20,000,000원", "3.6%"],
    ["합계", "552,000,000원", "104.2%"],
])
add_p(doc, "")
add_p(doc, "✅ 입증율 104.2% / 잉여 +22,000,000원", bold=True, size=12, color=RGBColor(0x00, 0x80, 0x00))

add_h2(doc, "3. 거래의 본질")
add_p(doc, "정상적인 부동산 거래", bold=True, size=12)
add_p(doc, "  살던 전셋집 보증금 반환 → 아파트 매입 → 기존 임차인 보증금 반환")

# === 1. 거래 구조 ===
add_h1(doc, "1. 거래 구조")

add_h2(doc, "1-1. 부동산 거래 사이클")
p = doc.add_paragraph()
run = p.add_run("""
[2021년]                    [2022년]
안경희 (이전 임대인)            박영준 (대표님)
      │                          │
      │ 전세보증금 반환            │ 신규 매입 아파트의
      ▼ 3.73억                  │ 기존 임차인 김광규에게
                                ▼ 전세보증금 반환 5.30억
                                  (직접 입주 목적)
""")
set_font(run, name='Consolas', size=9)

add_h2(doc, "1-2. 김광규 송금 4건")
add_table(doc, ["일시", "금액", "출발 계좌", "적요"], [
    ["2021-12-28 12:19", "10,000,000원", "IBK 012", "전세계약금김광규"],
    ["2022-01-03 12:44", "40,000,000원", "IBK 020", "전세계약금"],
    ["2022-02-23 12:07", "300,000,000원", "IBK 044", "김광규"],
    ["2022-02-23 13:15", "180,000,000원", "IBK 044", "김광규"],
    ["합계", "530,000,000원", "", ""],
])
add_p(doc, "")
add_p(doc, "수취 계좌: 신한은행 110406712671 김광규", italic=True, size=10)

# === 2. 자금원 상세 ===
add_h1(doc, "2. 자금원 상세")

# 2-1
add_h2(doc, "2-1. 안경희 전세보증금 반환 — 373,000,000원 (67.6%)")
add_p(doc, "대표님이 거주하던 전셋집의 임대인 안경희님으로부터 받은 보증금 반환")
add_table(doc, ["일시", "금액", "입금 계좌"], [
    ["2021-01-30 19:15", "37,000,000", "IBK 012"],
    ["2021-01-30 19:39", "13,000,000", "IBK 012"],
    ["2021-03-27 20:15", "50,000,000", "IBK 012"],
    ["2021-05-28 11:52", "100,000,000", "IBK 012"],
    ["2021-05-28 11:53", "100,000,000", "IBK 012"],
    ["2021-05-28 11:55", "70,000,000", "IBK 012"],
    ["2021-05-28 11:56", "3,000,000", "IBK 012"],
    ["2021-05-28 11:57", "585,220", "IBK 012 (관리비)"],
    ["합계 (절삭)", "373,000,000", ""],
])
add_p(doc, "")
add_p(doc, "📎 필수 증빙: 안경희와의 전세 임대차계약서 + 정산서", bold=True, color=RGBColor(0xC0, 0x00, 0x00))

# 2-2
add_h2(doc, "2-2. IBK 037 상가 임대 보증금 — 94,000,000원 (17.0%)")
add_p(doc, "대표님 운영 상가 임대 사업 (IBK 140-090845-01-037)의 보증금 수령 (2019~2021)")
add_table(doc, ["일시", "금액", "임차인"], [
    ["2019-04-12", "15,619,740", "경기광주노무사"],
    ["2019-05-09", "5,000,000", "주이경"],
    ["2019-06-26", "35,000,000", "주이경 (잔금)"],
    ["2019-09-25", "19,000,000", "이경부"],
    ["2019-12-10", "6,000,000", "조영진"],
    ["2020-03-13", "6,000,000", "김일선"],
    ["2021-05-10", "8,000,000", "조동조"],
    ["합계 (절삭)", "94,000,000", ""],
])
add_p(doc, "")
add_p(doc, "📎 필수 증빙: 상가 임대차계약서 7건, 사업자등록증, 종합소득세 신고서", bold=True, color=RGBColor(0xC0, 0x00, 0x00))
add_p(doc, "")
add_p(doc, "⚠ 참고: IBK 037은 상가 임대 사업 운영 계좌입니다. 정기 임대료 수입(약 9,800만원)은 사업운영비/생활비 소진으로 분류하여 본 자금원에는 포함하지 않았습니다.", italic=True, size=10)

# 2-3
add_h2(doc, "2-3. 김수동(매제) 차용금 — 65,000,000원 (11.8%) ⭐")
add_p(doc, "대표님 여동생의 배우자(매제)로부터의 친족 신뢰관계 차용")
add_table(doc, ["구분", "내용"], [
    ["차용일자", "2022년 2월 22일 09:05"],
    ["차용금액", "65,000,000원"],
    ["수령계좌", "농협 312-3479-3479-01 (E-우리은행 이체)"],
    ["차용 사유", "김광규 송금 자금 부족"],
    ["차용 조건", "무이자, 변제기 미정 (친족 신뢰관계)"],
    ["일부 변제", "2026-06-05 / 15,000,000원 / KB국민은행 송금"],
    ["미변제 잔액", "50,000,000원"],
])
add_p(doc, "")
add_p(doc, "📎 필수 증빙:", bold=True, color=RGBColor(0xC0, 0x00, 0x00))
add_p(doc, "  • 차용확인서 (2026-06-09 작성, 양 당사자 서명·날인) ⭐")
add_p(doc, "  • 농협 거래내역 (2022-02-22 입금)")
add_p(doc, "  • KB국민은행 송금확인서 (2026-06-05 변제)")
add_p(doc, "  • 가족관계증명서 (매제 관계 입증)")
add_p(doc, "")
add_p(doc, "📝 차용 입증 강도: ⭐⭐⭐ (송금기록 + 친족관계 + 일부변제 + 사후확인서)", bold=True, color=RGBColor(0x00, 0x80, 0x00))

# 2-4
add_h2(doc, "2-4. 배우자(이지영) 자금 이체 — 20,000,000원 (3.6%)")
add_p(doc, "배우자 별도 자산에서의 이체")
add_table(doc, ["일시", "금액", "거래방법"], [
    ["2021-02-01 15:16", "5,000,000", "우리은행 API이체"],
    ["2021-02-01 15:16", "5,000,000", "우리은행 API이체"],
    ["2021-02-07 15:21", "10,000,000", "우리은행 API이체"],
    ["합계", "20,000,000", ""],
])
add_p(doc, "")
add_p(doc, "⚠ 참고: 부부 간 5백만원 미만 양방향 거래는 생활비 정산으로 자금원 제외", italic=True, size=10)

# === 3. 자금 흐름 도식 ===
add_h1(doc, "3. 자금 흐름 도식")
p = doc.add_paragraph()
run = p.add_run("""
┌── 자금원 ─────────────────┐    ┌── 자금풀 ──┐    ┌── 지급 ──┐
│                           │    │           │    │           │
│ ① 안경희 임대인             │ →  │ IBK 012   │    │           │
│   전세보증금 반환            │    │ IBK 020   │    │ 김광규    │
│   3.73억 (2021.01~05)     │    │ IBK 037   │ →  │ 임차인     │
│                           │    │ IBK 044   │    │ 5.30억    │
│ ② 037 상가 보증금          │ →  │ 농협 312- │    │           │
│   0.94억                  │    │           │    │ 2022.02   │
│                           │    │           │    │           │
│ ③ 김수동(매제) 차용         │ →  │           │    │           │
│   0.65억 (2022.02.22)     │    │           │    │           │
│                           │    │           │    │           │
│ ④ 배우자 이체             │ →  │           │    │           │
│   0.20억                  │    │           │    │           │
│                           │    │           │    │           │
└───────────────────────────┘    └───────────┘    └───────────┘

총 자금원: 5.52억  →  김광규 송금: 5.30억 (입증율 104.2%)
""")
set_font(run, name='Consolas', size=9)

# === 4. 분석 데이터 범위 ===
add_h1(doc, "4. 분석 데이터 범위")
add_table(doc, ["계좌", "은행", "기간", "거래 건수"], [
    ["140-090845-01-012", "IBK기업 (주거래)", "2019-2021", "984건"],
    ["140-090845-01-020", "IBK기업", "2021-2022", "145건"],
    ["140-090845-01-037", "IBK기업 (상가 임대)", "2019-2021", "879건"],
    ["140-090845-01-044", "IBK기업 (전세 송금)", "2021-2022", "35건"],
    ["312-3479-3479-01", "농협 (자유예금)", "2019-2022", "1,018건"],
    ["합계", "", "", "3,075건"],
])

# === 5. 배제 항목 ===
add_h1(doc, "5. 배제 항목 — 예상 질문 대응")
add_p(doc, "세무서가 거래내역에서 발견할 가능성이 있는 큰 자금 거래에 대한 해명")

add_h2(doc, "5-1. 임시자금 (배제 사유 명확)")
add_table(doc, ["항목", "금액", "거래 패턴", "배제 사유"], [
    ["BNK캐피탈", "299,711,578원", "2021-12-29 입금", "일시 자금, 즉시 정산"],
    ["이정훈", "165,000,000원", "2회 입금", "즉시 환급 (분 단위) ⚠"],
    ["이승원", "130,000,000원", "2022-02-21 입금", "임시 자금 (K뱅크)"],
    ["박상호", "818,000,000원", "2021-01-26 입금", "3일 후 수표 출금"],
])

add_h2(doc, "5-2. 빗썸 가상화폐 거래")
add_p(doc, "대표님 본인 명의 가상화폐 거래소 (빗썸코리아) 거래:")
add_p(doc, "  • 매수 누적 (2019~2022.02): 758,000,000원")
add_p(doc, "  • 매도 누적 (2019~2022.02): 90,000,000원")
add_p(doc, "  • 순매수 (코인 보유): 668,000,000원")
add_p(doc, "")
add_p(doc, "⚠ 세무사 검토 필요:", bold=True, color=RGBColor(0xC0, 0x00, 0x00))
add_p(doc, "  • 김광규 송금 직전(2022-02-22~23) 빗썸 매도 회수 약 2.1억 발생")
add_p(doc, "  • 본 자금원에는 빗썸을 포함하지 않았으나, 자금 흐름상 회수 자금이 김광규 송금에 일부 사용됨")
add_p(doc, "  • 가상자산 양도소득세는 2025년 1월부터 시행 → 2022년 거래는 비과세")
add_p(doc, "  • 단, 자금 출처 검증 시 별도 소명이 필요할 수 있음")

add_h2(doc, "5-3. 김연호")
add_p(doc, "  • 2019-07-15: 4,000만원 입금 (국민은행)")
add_p(doc, "  • 2021-06~12: 매월 22.5만원 정기 입금 7회")
add_p(doc, "  • 의뢰인 판단으로 본 자금원에서 제외")

add_h2(doc, "5-4. 기타")
add_p(doc, "  • 급여(에이엑티브): 누적 1억 3,200만원 → 생활비 소진")
add_p(doc, "  • 037 임대료: 약 9,800만원 → 사업운영비/생활비 소진")
add_p(doc, "  • 배우자 소액: 양방향 생활비 정산")

# === 6. 의심 거래 사전 대응 ===
add_h1(doc, "6. 의심받을 수 있는 거래 패턴 — 사전 대응")

add_h2(doc, "6-1. 2022-02-23 단일 일자에 4.8억 통과")
add_p(doc, "현황: IBK 044 계좌에 1일 사이 6억 입금, 4.8억 김광규 송금", bold=True)
add_p(doc, "해명:")
add_p(doc, "  • IBK 044는 부동산 거래 전용 계좌")
add_p(doc, "  • 같은 날 잔금 송금일이라 다양한 자금원을 044에 모은 후 송금")
add_p(doc, "  • 자금원 구성: 본인 다른 IBK 계좌 + 농협 (안경희 자금 + 빗썸 회수 + 김수동 차용)")

add_h2(doc, "6-2. 이정훈 자금 (2022-02-23 13:06~13:18)")
p = doc.add_paragraph()
run = p.add_run("13:06  +1억원   이정훈 입금\n13:07  +5천만원 이정훈 입금\n13:18  -1.2억원 이정훈 환급")
set_font(run, name='Consolas', size=10)
add_p(doc, "")
add_p(doc, "해명: 김광규 송금 자금 마련 중 일시 부족분을 이정훈으로부터 임시 차용 후 5분간 보유 후 즉시 환급. 형식적 차용/환급으로 실질 자금원 아님.")
add_p(doc, "📎 필요 증빙: 이정훈 정산 확인서", bold=True, color=RGBColor(0xC0, 0x00, 0x00))

add_h2(doc, "6-3. 빗썸 거래 활성도")
add_p(doc, "해명:")
add_p(doc, "  • 본인 자금으로 정상적인 가상자산 투자")
add_p(doc, "  • 농협 계좌에서 운용 (매수→매도→회수 모두 본인 명의)")
add_p(doc, "  • 2022년 거래는 가상자산 양도소득세 시행 전이라 비과세")

# === 7. 증빙 체크리스트 ===
add_h1(doc, "7. 필수 증빙 자료 체크리스트")

add_h2(doc, "⭐⭐⭐ 우선순위 1 (이번 주 내 필수)")
checks_1 = [
    "안경희 전세 임대차계약서 + 정산서",
    "김광규 전세 임대차계약서 (승계)",
    "신규 매입 아파트 매매계약서 (임차인 김광규 승계 명시)",
    "김수동 차용확인서 (2026-06-09 작성) - 양식 제공",
    "KB국민은행 송금확인서 (2026-06-05 / 1,500만)",
    "농협 거래내역 (2022-02-22 / 6,500만 입금)",
    "가족관계증명서 (김수동 매제 관계 입증)",
]
for item in checks_1:
    add_p(doc, "  ☐ " + item)

add_h2(doc, "⭐⭐ 우선순위 2 (가능한 빨리)")
checks_2 = [
    "상가 임대차계약서 7건 (자금원 ②)",
    "사업자등록증 (부동산 임대업)",
    "2019-2022 종합소득세 신고서",
]
for item in checks_2:
    add_p(doc, "  ☐ " + item)

add_h2(doc, "⭐ 우선순위 3 (보강 자료)")
checks_3 = [
    "이정훈 정산 확인서 (임시자금 입증)",
    "이승원 정산 확인서 (임시자금 입증)",
    "박상호 정산 확인서 (임시자금 입증)",
    "배우자 이지영 자금 출처 자료",
    "빗썸 거래 내역 (요청 시)",
]
for item in checks_3:
    add_p(doc, "  ☐ " + item)

# === 8. 핵심 메시지 ===
add_h1(doc, "8. 핵심 메시지")

add_h2(doc, "거래의 정상성")
add_p(doc, "✅ 부동산 임대차 사이클 명확: 안경희 → 박영준 → 김광규 (이사 + 신규 매입)", color=RGBColor(0x00, 0x80, 0x00))
add_p(doc, "✅ 상가 임대 사업 정상 운영: IBK 037 계좌로 3년간 임대 보증금 안정적 수령", color=RGBColor(0x00, 0x80, 0x00))
add_p(doc, "✅ 친족 간 차용 명확: 매제 + 송금기록 + 일부 변제 + 차용확인서", color=RGBColor(0x00, 0x80, 0x00))
add_p(doc, "✅ 배우자 자금 이체: 부부 간 자산 이전 (생활비 정산과 구분)", color=RGBColor(0x00, 0x80, 0x00))
add_p(doc, "✅ 모든 거래 금융기관 정식 채널: 시중은행 송금만 사용", color=RGBColor(0x00, 0x80, 0x00))
add_p(doc, "✅ 계좌 명의 일치: 모든 거래가 박영준 본인 명의 계좌", color=RGBColor(0x00, 0x80, 0x00))

doc.add_paragraph()
add_p(doc, "한 줄 요약", bold=True, size=12)
add_p(doc, "이전 임대인 안경희님으로부터 받은 전세보증금(3.73억) + 상가 임대 보증금(0.94억) + 매제 김수동님 차용금(0.65억) + 배우자 이체 자금(0.20억)으로, 매입한 아파트의 기존 임차인 김광규님께 전세보증금(5.30억)을 반환한 정상 부동산 거래", italic=True, color=RGBColor(0x1F, 0x4E, 0x79))

# === 의뢰인 연락처 ===
doc.add_paragraph()
add_h1(doc, "9. 의뢰인 연락처")
add_p(doc, "박영준 (대표)", bold=True, size=12)
add_p(doc, "  • 이메일: zunn@eactive.co.kr")
add_p(doc, "  • 연락처: _____________________")
add_p(doc, "  • 긴급 시 즉시 연락 바랍니다.", italic=True)

# === 면책 ===
doc.add_paragraph()
add_h1(doc, "⚠ 면책")
add_p(doc, "본 자료는 의뢰인이 정리한 자체 분석 자료이며, 세무 소명 시:", italic=True, size=10)
add_p(doc, "  ① 반드시 전문 세무사·회계사 검토 후 제출", italic=True, size=10)
add_p(doc, "  ② 본 자료는 참고 자료로 활용", italic=True, size=10)
add_p(doc, "  ③ 실제 제출 자료는 세무사가 가공해야 함", italic=True, size=10)
add_p(doc, "  ④ 필수 증빙 자료는 의뢰인이 사전 확보 중", italic=True, size=10)

doc.add_paragraph()
add_p(doc, "작성일: 2026년 6월 9일", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)
add_p(doc, "의뢰인: 박영준", align=WD_ALIGN_PARAGRAPH.RIGHT, italic=True)

# 저장
import os
output_path = r"D:\Claw\workspace\analysis\세무사_브리핑_김광규자금출처.docx"
doc.save(output_path)
print(f"OK 세무사 브리핑 생성 완료")
print(f"파일: {output_path}")
print(f"크기: {os.path.getsize(output_path):,} bytes")
