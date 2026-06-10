# -*- coding: utf-8 -*-
"""차용확인서 DOCX 생성 (한글 폰트 안정 버전)"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name='맑은 고딕', size=11, bold=False, italic=False):
    """한글 폰트 안전 설정"""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)


def add_p(doc, text, size=11, bold=False, align=None, italic=False):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, size=level, bold=True)
    return p


def make_kv_table(doc, data, col_widths=(4, 11)):
    table = doc.add_table(rows=len(data), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = k
        row.cells[1].text = v
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_font(r, size=11)
    return table


doc = Document()

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

# === 제목 ===
add_p(doc, "차 용 확 인 서", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p(doc, "문서 작성일: 2026년 6월 9일", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()

# === 1. 당사자 ===
add_heading(doc, "1. 당사자")
add_p(doc, "【채무자 (借用人)】", bold=True)
make_kv_table(doc, [
    ("성        명", "박 영 준"),
    ("주민등록번호", "                  -                  "),
    ("주        소", ""),
    ("연  락  처", ""),
])
doc.add_paragraph()
add_p(doc, "【채권자 (貸與人)】", bold=True)
make_kv_table(doc, [
    ("성        명", "김 수 동"),
    ("주민등록번호", "                  -                  "),
    ("주        소", ""),
    ("연  락  처", ""),
    ("채무자와의 관계", "처남 (채무자 박영준의 여동생의 배우자)"),
])
doc.add_paragraph()

# === 2. 차용 사실 ===
add_heading(doc, "2. 차용 사실 확인")
add_p(doc, "위 채무자(박영준)와 채권자(김수동)는 아래와 같이 금원의 차용 사실을 확인합니다.")
add_p(doc, "2.1 차용 내역", bold=True)
make_kv_table(doc, [
    ("차용 일자", "2022년 2월 22일"),
    ("차용 금액", "금 육천오백만원 (₩65,000,000)"),
    ("송금 방법", "은행 계좌 이체 (E-우리은행)"),
    ("수령 계좌", "농협은행 312-3479-3479-01 (예금주: 박영준)"),
    ("송금 시각", "2022년 2월 22일 09시 05분"),
    ("차용 목적", "채무자(박영준)의 부동산 거래 자금 (전세 보증금 반환)"),
])
doc.add_paragraph()

add_p(doc, "2.2 차용 조건", bold=True)
make_kv_table(doc, [
    ("이 자 율", "무이자 (친족 신뢰관계)"),
    ("변제 기일", "정함이 없음 (채무자 자금 사정에 따라 분할 변제)"),
    ("변제 방법", "채무자 변제 능력 회복 시 분할 상환"),
])

# === 3. 차용증 미작성 경위 ===
doc.add_paragraph()
add_heading(doc, "3. 차용증 미작성 경위")
add_p(doc, "본 차용 당시 정식 차용증을 작성하지 아니한 사유는 다음과 같습니다.")
add_p(doc, "① 친족 관계: 채권자(김수동)는 채무자(박영준)의 여동생의 배우자(매제)로서, 가족 간 신뢰관계에 기초한 자금 거래임.")
add_p(doc, "② 사회 통념: 가족·친족 간 자금 거래에서는 정식 차용증 작성 없이 신뢰관계로 금원을 차용하는 것이 사회 통념상 일반적임.")
add_p(doc, "③ 차용 목적의 명확성: 차용 자금이 채무자의 부동산 거래(전세 보증금 반환) 용도로 명확히 사용되었으며, 양 당사자가 차용 사실을 명확히 인식하고 있었음.")
add_p(doc, "따라서 본 확인서는 차용 사실 자체를 부인하는 것이 아니라, 사후적으로 차용 조건을 명문화하기 위한 것입니다.")

# === 4. 변제 사실 ===
doc.add_paragraph()
add_heading(doc, "4. 변제 사실")

add_p(doc, "4.1 변제 능력 부족 기간 (2022년 2월 ~ 2026년 5월)", bold=True)
add_p(doc, "채무자(박영준)는 차용 이후 다음과 같은 사유로 즉시 변제가 어려웠음을 확인합니다.")
add_p(doc, "  • 부동산 거래 자금 운용으로 가용 자금 부족")
add_p(doc, "  • 상가 임대 사업 운영비 부담")
add_p(doc, "  • 변제 능력 회복 시점까지 채권자(김수동)와 변제 시기를 협의")

add_p(doc, "4.2 일부 변제 사실", bold=True)
table = doc.add_table(rows=2, cols=4)
table.style = 'Light Grid Accent 1'
headers = ["변제 회차", "변제 일자", "변제 금액", "변제 방법"]
values = ["1차", "2026년 6월 5일", "금 일천오백만원 (₩15,000,000)", "KB국민은행 송금"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for r in table.rows[0].cells[i].paragraphs[0].runs:
        set_font(r, bold=True, size=11)
for i, v in enumerate(values):
    table.rows[1].cells[i].text = v
    for r in table.rows[1].cells[i].paragraphs[0].runs:
        set_font(r, size=11)

doc.add_paragraph()
add_p(doc, "4.3 변제 사유 (2026년 6월 5일)", bold=True)
add_p(doc, "  • 채권자(김수동) 측 자금 필요로 변제 요청")
add_p(doc, "  • 채무자(박영준) 자금 사정 호전으로 일부 변제 가능")
add_p(doc, "  • 양 당사자 합의에 따라 우선 1,500만원 변제")

add_p(doc, "4.4 미변제 잔액", bold=True)
p = doc.add_paragraph()
run = p.add_run("원래 차용금:            ₩65,000,000\n")
set_font(run, name='Consolas', size=11)
run = p.add_run("1차 변제 (2026-06-05):  -₩15,000,000\n")
set_font(run, name='Consolas', size=11)
run = p.add_run("─────────────────────────────────\n")
set_font(run, name='Consolas', size=11)
run = p.add_run("미변제 잔액:            ₩50,000,000")
set_font(run, name='Consolas', size=11, bold=True)

add_p(doc, "현재 미변제 잔액: 금 오천만원 (₩50,000,000)", bold=True, size=12)

add_p(doc, "4.5 향후 변제 계획", bold=True)
add_p(doc, "  • 채무자의 자금 사정에 따라 분할 변제 예정")
add_p(doc, "  • 양 당사자가 협의하여 변제 일정 조정")

# === 5. 확인서 작성 사유 ===
doc.add_paragraph()
add_heading(doc, "5. 확인서 작성 사유")
add_p(doc, "본 확인서는 다음과 같은 사유로 2026년 6월 9일 작성되었습니다.")
add_p(doc, "① 채무자(박영준)의 자금 출처 정리 및 세무 소명 자료 준비")
add_p(doc, "② 양 당사자 간 차용 사실 및 변제 내역의 객관적 확인")
add_p(doc, "③ 향후 변제 진행 시 기준 자료로 활용")
add_p(doc, "본 확인서는 차용 사실을 사후에 인정·확인하는 문서이며, 차용 일자(2022년 2월 22일)에 작성된 것이 아님을 명백히 밝힙니다.")

# === 6. 첨부 자료 ===
doc.add_paragraph()
add_heading(doc, "6. 첨부 자료")
add_p(doc, "□ 첨부 1: 농협은행 거래내역 (2022년 2월 22일 65,000,000원 입금)")
add_p(doc, "□ 첨부 2: KB국민은행 송금 확인서 (2026년 6월 5일 15,000,000원 송금)")
add_p(doc, "□ 첨부 3: 가족관계증명서 (채무자와 채권자 간 친족 관계 확인)")
add_p(doc, "□ 첨부 4: 채무자 신분증 사본")
add_p(doc, "□ 첨부 5: 채권자 신분증 사본")

# === 7. 확인 및 서명 ===
doc.add_paragraph()
add_heading(doc, "7. 확인 및 서명")
add_p(doc, "위 사실에 대하여 채무자(박영준)와 채권자(김수동)는 모두 동의하며, 본 확인서의 내용이 사실과 일치함을 확인합니다.")
doc.add_paragraph()
add_p(doc, "작 성 일 :  2 0 2 6 년   6 월   9 일", bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

# 서명란 (간단 형식)
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# 채무자
cell = table.rows[0].cells[0]
cell.text = ""
sign_lines = [
    ("【채무자】", True, 12),
    ("", False, 11),
    ("", False, 11),
    ("성  명:  박 영 준          (인)", False, 11),
    ("", False, 11),
    ("주민등록번호:", False, 11),
    ("                  -", False, 11),
    ("", False, 11),
    ("주    소:", False, 11),
    ("", False, 11),
    ("", False, 11),
]
for text, bold, sz in sign_lines:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=sz, bold=bold)

# 채권자
cell = table.rows[0].cells[1]
cell.text = ""
sign_lines2 = [
    ("【채권자】", True, 12),
    ("", False, 11),
    ("", False, 11),
    ("성  명:  김 수 동          (인)", False, 11),
    ("", False, 11),
    ("주민등록번호:", False, 11),
    ("                  -", False, 11),
    ("", False, 11),
    ("주    소:", False, 11),
    ("", False, 11),
    ("", False, 11),
]
for text, bold, sz in sign_lines2:
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=sz, bold=bold)

# 저장
import os
output_path = r"D:\Claw\workspace\analysis\차용확인서_김수동_박영준.docx"
doc.save(output_path)
print(f"OK 차용확인서 생성 완료")
print(f"파일: {output_path}")
print(f"크기: {os.path.getsize(output_path):,} bytes")
